"""
AI Business Strategy & Competitor Intelligence Agent
-----------------------------------------------------
Streamlit app that:
  1. Discovers competitors for a given business domain/region (DDGS web search,
     using several query variants merged into a unique competitor list)
  2. Scrapes competitor websites — homepage + common subpages like /about,
     /pricing, /features (requests + BeautifulSoup)
  3. Uses Groq for ALL repetitive extraction/analysis (competitor profiles,
     market SWOT, business-consultant chat) and reserves Google Gemini for the
     one task that benefits most from deep reasoning: the final growth
     strategy + 30/60/90-day plan. This keeps Gemini call volume low and
     avoids Gemini free-tier rate-limit/quota errors.
  4. Generates a downloadable, professional .docx business intelligence report

Model responsibility map:
    Groq   -> analyze_competitor(), build_swot(), answer_business_question()
    Gemini -> build_growth_strategy() only

Run with:
    streamlit run app.py

Requires a .env file (see .env.example) with:
    GROQ_API_KEY=...
    GEMINI_API_KEY=...
"""

import os
import json
import time
import textwrap
from datetime import datetime
from urllib.parse import urlparse

import requests
import streamlit as st
from bs4 import BeautifulSoup
from dotenv import load_dotenv

# ---- Optional / third-party clients -----------------------------------------
try:
    from ddgs import DDGS  # pip install ddgs
except ImportError:  # pragma: no cover
    DDGS = None

try:
    from groq import Groq
except ImportError:  # pragma: no cover
    Groq = None

try:
    import google.generativeai as genai
except ImportError:  # pragma: no cover
    genai = None

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

# ==============================================================================
# CONFIG
# ==============================================================================

st.set_page_config(
    page_title="AI Business Strategy & Competitor Intelligence Agent",
    page_icon="🚀",
    layout="wide",
)

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-flash")

REQUEST_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
    )
}


# ==============================================================================
# LLM CLIENT WRAPPERS
# ==============================================================================

class LLMClients:
    """Small wrapper so the rest of the app doesn't care which provider is used."""

    def __init__(self, groq_key: str, gemini_key: str):
        self.groq_client = None
        self.gemini_model = None

        if groq_key and Groq is not None:
            try:
                self.groq_client = Groq(api_key=groq_key)
            except Exception as e:  # pragma: no cover
                st.warning(f"Could not initialize Groq client: {e}")

        if gemini_key and genai is not None:
            try:
                genai.configure(api_key=gemini_key)
                self.gemini_model = genai.GenerativeModel(GEMINI_MODEL)
            except Exception as e:  # pragma: no cover
                st.warning(f"Could not initialize Gemini client: {e}")

    def is_ready(self) -> bool:
        return self.groq_client is not None or self.gemini_model is not None

    def fast_extract(self, prompt: str) -> str:
        """
        Use Groq for quick, structured extraction tasks. This is now the
        workhorse for nearly everything (competitor profiles, SWOT, chat) so
        that Gemini's quota is reserved for build_growth_strategy() only.

        NOTE: this intentionally does NOT fall back to Gemini on failure —
        that silent fallback was a major source of unexpected Gemini calls
        (and quota errors) in the previous version. If Groq fails, we return
        "" and let the caller's fallback/default data kick in instead.
        """
        if self.groq_client is None:
            st.warning("Groq is not configured — skipping this analysis step.")
            return ""
        try:
            resp = self.groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=1500,
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            msg = str(e).lower()
            if "rate" in msg or "quota" in msg or "429" in msg:
                st.warning("Groq rate limit/quota hit — continuing with partial data.")
            else:
                st.warning(f"Groq call failed: {e}")
            return ""

    def deep_reason(self, prompt: str) -> str:
        """
        Use Gemini for deeper strategic reasoning. Reserved for
        build_growth_strategy() so we stay well within free-tier quotas.
        """
        if self.gemini_model is None:
            st.warning("Gemini is not configured — skipping deep strategic reasoning step.")
            return ""
        try:
            resp = self.gemini_model.generate_content(prompt)
            return (resp.text or "").strip()
        except Exception as e:
            msg = str(e).lower()
            if "quota" in msg or "rate" in msg or "429" in msg or "resourceexhausted" in msg:
                st.warning(
                    "Gemini quota/rate limit reached. The growth strategy section "
                    "may be incomplete — try again in a minute, or rely on the "
                    "Groq-generated sections in the meantime."
                )
            else:
                st.warning(f"Gemini call failed: {e}")
            return ""

    def ask(self, prompt: str, prefer: str = "fast") -> str:
        if prefer == "fast":
            return self.fast_extract(prompt)
        return self.deep_reason(prompt)


def safe_json_parse(text: str, fallback: dict) -> tuple:
    """
    LLMs sometimes wrap JSON in markdown fences or add preamble; clean it up.
    Returns (parsed_dict_or_fallback, success_bool) so callers can decide
    whether to retry.
    """
    if not text:
        return fallback, False
    cleaned = text.strip()
    cleaned = cleaned.replace("```json", "").replace("```", "").strip()
    # Try to find the outermost braces if there's stray text around the JSON
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start != -1 and end != -1 and end > start:
        cleaned = cleaned[start : end + 1]
    try:
        return json.loads(cleaned), True
    except Exception:
        return fallback, False


def call_llm_for_json(llm_call_fn, prompt: str, fallback: dict, retries: int = 1) -> dict:
    """
    Calls an LLM function (llm.fast_extract or llm.deep_reason), parses the
    JSON response, and retries once with a stricter instruction if parsing
    fails. Falls back gracefully (never raises) if all attempts fail.
    """
    attempt_prompt = prompt
    for attempt in range(retries + 1):
        raw = llm_call_fn(attempt_prompt)
        parsed, ok = safe_json_parse(raw, fallback)
        if ok:
            return parsed
        if attempt < retries:
            attempt_prompt = (
                prompt
                + "\n\nIMPORTANT: Your previous response was not valid JSON. "
                "Respond with ONLY a single valid JSON object — no markdown "
                "fences, no commentary, no explanation."
            )
    # All attempts failed — continue the app with the fallback instead of crashing
    return fallback


# ==============================================================================
# WEB SEARCH + SCRAPING
# ==============================================================================

BAD_DOMAINS = [
    "wikipedia.org", "linkedin.com", "youtube.com", "reddit.com",
    "facebook.com", "twitter.com", "x.com", "instagram.com",
    "medium.com", "crunchbase.com", "g2.com", "capterra.com",
    "quora.com", "pinterest.com",
]


def build_search_queries(business_domain: str, region: str) -> list:
    """Several phrasings tend to surface a broader, less redundant set of
    competitors than a single query."""
    return [
        f"top {business_domain} companies in {region}",
        f"best {business_domain} in {region}",
        f"leading {business_domain} companies {region}",
        f"{business_domain} competitors {region}",
    ]


def discover_competitors(business_domain: str, region: str, max_results: int = 6) -> list:
    """Use DDGS web search (multiple query variants) to find likely competitor
    websites, merging unique domains across all queries."""
    if DDGS is None:
        st.error("The `ddgs` package is not installed. Run: pip install ddgs")
        return []

    queries = build_search_queries(business_domain, region)
    results = []
    seen_domains = set()

    try:
        with DDGS() as ddgs:
            for query in queries:
                if len(results) >= max_results:
                    break
                try:
                    hits = ddgs.text(query, max_results=max_results * 2)
                except Exception as e:
                    # One bad query shouldn't kill the whole discovery step
                    st.warning(f"Search failed for query '{query}': {e}")
                    continue

                for r in hits:
                    if len(results) >= max_results:
                        break
                    url = r.get("href") or r.get("url")
                    title = r.get("title", "")
                    if not url:
                        continue
                    domain = urlparse(url).netloc.replace("www.", "")
                    if not domain or domain in seen_domains:
                        continue
                    if any(bad in domain for bad in BAD_DOMAINS):
                        continue
                    seen_domains.add(domain)
                    results.append({"name": title, "domain": domain, "url": url})
    except Exception as e:
        st.error(f"Competitor search failed: {e}")

    return results


COMMON_SUBPAGES = ["/about", "/services", "/products", "/pricing", "/features", "/solutions"]

# Boilerplate snippets that add noise/tokens without business signal
NOISE_PHRASES = [
    "accept cookies", "we use cookies", "cookie policy", "cookie preferences",
    "all rights reserved", "privacy policy", "terms of service", "terms of use",
    "subscribe to our newsletter", "sign up for our newsletter",
]


def _clean_html_to_text(html: str) -> str:
    """Strip nav/scripts/styles and repetitive boilerplate, keep business content."""
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "header", "footer", "nav",
                      "form", "iframe", "button"]):
        tag.decompose()

    # Drop elements that look like cookie banners / consent widgets
    for tag in soup.find_all(True):
        classes_and_ids = " ".join(tag.get("class", []) + [tag.get("id", "")]).lower()
        if any(kw in classes_and_ids for kw in ["cookie", "consent", "gdpr", "popup", "modal"]):
            tag.decompose()

    text = " ".join(soup.get_text(separator=" ").split())

    # Remove sentences that are pure boilerplate noise
    for phrase in NOISE_PHRASES:
        text = text.replace(phrase, "")

    return " ".join(text.split())


def scrape_website(url: str, max_chars: int = 2000) -> str:
    """Fetch and clean a single page's text. Returns '' on failure instead of
    an error string, so callers can decide how to handle missing content."""
    try:
        resp = requests.get(url, headers=REQUEST_HEADERS, timeout=10)
        resp.raise_for_status()
    except Exception:
        return ""
    return _clean_html_to_text(resp.text)[:max_chars]


def scrape_competitor_full(base_url: str, per_page_chars: int = 1200, total_max_chars: int = 3000) -> str:
    """
    Scrape the homepage plus common subpages (/about, /pricing, /features, ...)
    when they exist, and merge the text — giving the LLM a fuller picture of
    the business than the homepage alone, while still keeping token usage low
    via a total character cap.
    """
    base_url = base_url.rstrip("/")
    combined_parts = []

    # Homepage first — most important page, gets a slightly larger slice
    homepage_text = scrape_website(base_url, max_chars=per_page_chars + 500)
    if homepage_text:
        combined_parts.append(homepage_text)

    for path in COMMON_SUBPAGES:
        if len("".join(combined_parts)) >= total_max_chars:
            break
        try:
            page_text = scrape_website(base_url + path, max_chars=per_page_chars)
        except Exception:
            page_text = ""
        if page_text:
            combined_parts.append(f"[{path}] {page_text}")
        time.sleep(0.1)  # be polite to the target server

    if not combined_parts:
        return f"[No content could be scraped from {base_url}]"

    merged = " ".join(combined_parts)
    return merged[:total_max_chars]


# ==============================================================================
# ANALYSIS PIPELINE
# ==============================================================================

COMPETITOR_EXTRACT_PROMPT = textwrap.dedent(
    """
    You are a business intelligence analyst. Based on the raw website text below
    for the company "{name}" ({url}), extract structured information.

    Respond with ONLY valid JSON, no preamble, matching this schema:
    {{
      "company_overview": "2-3 sentence summary of what the company does",
      "products_services": ["list", "of", "key products or services"],
      "pricing_info": "short summary of pricing model/signals, or 'Not publicly listed'",
      "target_audience": "who they appear to sell to",
      "marketing_strategy": "short summary of messaging, positioning, tone",
      "key_features": ["notable", "features", "or", "differentiators"],
      "calls_to_action": ["primary CTAs found on the site"],
      "strengths": ["2-4 strengths"],
      "weaknesses": ["2-4 weaknesses or gaps"]
    }}

    Raw website text:
    ---
    {content}
    ---
    """
)

SWOT_PROMPT = textwrap.dedent(
    """
    You are a senior strategy consultant. Using the competitor profiles below for
    the "{business_domain}" industry in "{region}", produce an overall market SWOT
    analysis (for a new/existing player considering this market), plus a ranked
    list of 3-5 concrete market gaps / opportunities.

    Respond with ONLY valid JSON matching this schema:
    {{
      "strengths": ["market-level strengths a player could leverage"],
      "weaknesses": ["market-level weaknesses/challenges"],
      "opportunities": ["specific, actionable opportunities"],
      "threats": ["market-level threats/risks"],
      "market_gaps": ["3-5 specific unmet needs or underserved segments"]
    }}

    Competitor profiles (JSON):
    {profiles}
    """
)

GROWTH_STRATEGY_PROMPT = textwrap.dedent(
    """
    You are an AI growth strategist. Given the competitor analysis and SWOT
    below for a business in "{business_domain}" targeting "{region}"
    (own website: {own_site}), produce a practical growth strategy.

    Respond with ONLY valid JSON matching this schema:
    {{
      "product_improvements": ["..."],
      "new_features": ["..."],
      "marketing_ideas": ["..."],
      "seo_recommendations": ["..."],
      "content_strategy": ["..."],
      "customer_acquisition_ideas": ["..."],
      "branding_suggestions": ["..."],
      "expansion_opportunities": ["..."],
      "plan_30_days": ["3-5 concrete actions"],
      "plan_60_days": ["3-5 concrete actions"],
      "plan_90_days": ["3-5 concrete actions"]
    }}

    Competitor profiles:
    {profiles}

    Market SWOT:
    {swot}
    """
)


def analyze_competitor(llm: LLMClients, name: str, url: str, content: str) -> dict:
    """Per-competitor extraction — Groq (fast, repetitive, high call volume).
    Unchanged from the original implementation, now using the retrying JSON helper."""
    prompt = COMPETITOR_EXTRACT_PROMPT.format(name=name, url=url, content=content)
    fallback = {
        "company_overview": "Not available",
        "products_services": [],
        "pricing_info": "Not available",
        "target_audience": "Not available",
        "marketing_strategy": "Not available",
        "key_features": [],
        "calls_to_action": [],
        "strengths": [],
        "weaknesses": [],
    }
    return call_llm_for_json(llm.fast_extract, prompt, fallback)


def build_swot(llm: LLMClients, business_domain: str, region: str, profiles: list) -> dict:
    """Market SWOT — moved to Groq (llm.fast_extract) so it no longer competes
    with the growth strategy step for Gemini's limited quota."""
    prompt = SWOT_PROMPT.format(
        business_domain=business_domain,
        region=region,
        profiles=json.dumps(profiles, indent=2)[:8000],
    )
    fallback = {
        "strengths": [], "weaknesses": [], "opportunities": [],
        "threats": [], "market_gaps": [],
    }
    return call_llm_for_json(llm.fast_extract, prompt, fallback)


def build_growth_strategy(llm: LLMClients, business_domain: str, region: str,
                           own_site: str, profiles: list, swot: dict) -> dict:
    """The ONLY step that uses Gemini (llm.deep_reason). This is the single
    highest-value place for deep strategic reasoning, and running it just
    once per analysis keeps us well within Gemini's rate limits."""
    prompt = GROWTH_STRATEGY_PROMPT.format(
        business_domain=business_domain,
        region=region,
        own_site=own_site or "Not provided",
        profiles=json.dumps(profiles, indent=2)[:6000],
        swot=json.dumps(swot, indent=2)[:3000],
    )
    fallback = {
        "product_improvements": [], "new_features": [], "marketing_ideas": [],
        "seo_recommendations": [], "content_strategy": [],
        "customer_acquisition_ideas": [], "branding_suggestions": [],
        "expansion_opportunities": [], "plan_30_days": [], "plan_60_days": [],
        "plan_90_days": [],
    }
    return call_llm_for_json(llm.deep_reason, prompt, fallback)


def answer_business_question(llm: LLMClients, question: str, context: dict) -> str:
    """Business consultant chat — moved to Groq (llm.fast_extract) so ad-hoc
    Q&A doesn't consume Gemini quota either."""
    prompt = f"""
    You are an AI business consultant. Using the research context below, answer
    the user's question clearly and specifically. Keep it concise but actionable.

    Context (competitor profiles, SWOT, growth strategy):
    {json.dumps(context, indent=2)[:8000]}

    Question: {question}
    """
    answer = llm.fast_extract(prompt)
    if not answer:
        answer = "I couldn't reach the AI service just now. Please try asking again in a moment."
    return answer


# ==============================================================================
# REPORT GENERATION (.docx)
# ==============================================================================

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_bullets(doc, items):
    for item in items or []:
        doc.add_paragraph(str(item), style="List Bullet")


def generate_report(business_domain, region, own_site, competitors, profiles, swot, strategy) -> str:
    doc = Document()

    title = doc.add_heading("Business Intelligence Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Industry: {business_domain}  |  Region: {region}  |  "
        f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    ).italic = True

    # Executive Summary
    add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph(
        f"This report analyzes {len(competitors)} key competitors in the "
        f"{business_domain} space within {region}. It summarizes each "
        "competitor's positioning, compares their offerings, and outlines a "
        "SWOT-based growth strategy with a 30/60/90-day action plan."
    )

    # Competitor Profiles
    add_heading(doc, "Competitor Profiles", level=1)
    for comp, profile in zip(competitors, profiles):
        add_heading(doc, f"{comp['name']} ({comp['domain']})", level=2)
        doc.add_paragraph(profile.get("company_overview", "Not available"))

        doc.add_paragraph("Products & Services:", style="Intense Quote")
        add_bullets(doc, profile.get("products_services"))

        doc.add_paragraph("Pricing: " + str(profile.get("pricing_info", "Not available")))
        doc.add_paragraph("Target Audience: " + str(profile.get("target_audience", "Not available")))
        doc.add_paragraph("Marketing Strategy: " + str(profile.get("marketing_strategy", "Not available")))

        doc.add_paragraph("Key Features:", style="Intense Quote")
        add_bullets(doc, profile.get("key_features"))

        doc.add_paragraph("Strengths:", style="Intense Quote")
        add_bullets(doc, profile.get("strengths"))

        doc.add_paragraph("Weaknesses:", style="Intense Quote")
        add_bullets(doc, profile.get("weaknesses"))

    # Comparison Table
    add_heading(doc, "Competitor Comparison", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Competitor", "Target Audience", "Pricing", "Key Strength"
    )
    for comp, profile in zip(competitors, profiles):
        row = table.add_row().cells
        row[0].text = comp["name"] or comp["domain"]
        row[1].text = str(profile.get("target_audience", "-"))
        row[2].text = str(profile.get("pricing_info", "-"))
        strengths = profile.get("strengths") or ["-"]
        row[3].text = strengths[0]

    # SWOT
    add_heading(doc, "Market SWOT Analysis", level=1)
    for label, key in [
        ("Strengths", "strengths"), ("Weaknesses", "weaknesses"),
        ("Opportunities", "opportunities"), ("Threats", "threats"),
    ]:
        doc.add_paragraph(label, style="Intense Quote")
        add_bullets(doc, swot.get(key))

    add_heading(doc, "Market Gaps & Opportunities", level=2)
    add_bullets(doc, swot.get("market_gaps"))

    # Growth Strategy
    add_heading(doc, "Growth Strategy Recommendations", level=1)
    sections = [
        ("Product Improvements", "product_improvements"),
        ("New Features", "new_features"),
        ("Marketing Ideas", "marketing_ideas"),
        ("SEO Recommendations", "seo_recommendations"),
        ("Content Strategy", "content_strategy"),
        ("Customer Acquisition Ideas", "customer_acquisition_ideas"),
        ("Branding Suggestions", "branding_suggestions"),
        ("Expansion Opportunities", "expansion_opportunities"),
    ]
    for label, key in sections:
        doc.add_paragraph(label, style="Intense Quote")
        add_bullets(doc, strategy.get(key))

    # Action Plan
    add_heading(doc, "30 / 60 / 90-Day Action Plan", level=1)
    for label, key in [
        ("First 30 Days", "plan_30_days"),
        ("Next 60 Days", "plan_60_days"),
        ("First 90 Days", "plan_90_days"),
    ]:
        add_heading(doc, label, level=2)
        add_bullets(doc, strategy.get(key))

    out_path = os.path.join(os.getcwd(), "business_intelligence_report.docx")
    doc.save(out_path)
    return out_path


# ==============================================================================
# STREAMLIT UI
# ==============================================================================

def init_session_state():
    for key, default in [
        ("competitors", []),
        ("profiles", []),
        ("swot", {}),
        ("strategy", {}),
        ("report_path", None),
        ("chat_history", []),
    ]:
        if key not in st.session_state:
            st.session_state[key] = default


def sidebar_inputs():
    st.sidebar.header("⚙️ Configuration")
    business_domain = st.sidebar.text_input("Business Domain / Industry", placeholder="e.g. project management software")
    region = st.sidebar.text_input("Target Region", placeholder="e.g. United States")
    own_site = st.sidebar.text_input("Your Website (optional)", placeholder="https://yourcompany.com")
    max_competitors = st.sidebar.slider("Number of competitors to analyze", 2, 8, 4)

    st.sidebar.markdown("---")
    groq_ok = "✅" if GROQ_API_KEY else "❌"
    gemini_ok = "✅" if GEMINI_API_KEY else "❌"
    st.sidebar.markdown(f"Groq API Key: {groq_ok}")
    st.sidebar.markdown(f"Gemini API Key: {gemini_ok}")
    if not GROQ_API_KEY and not GEMINI_API_KEY:
        st.sidebar.error("Add GROQ_API_KEY and/or GEMINI_API_KEY to your .env file.")

    run = st.sidebar.button("🚀 Run Competitive Analysis", type="primary", use_container_width=True)
    return business_domain, region, own_site, max_competitors, run


def run_pipeline(llm, business_domain, region, own_site, max_competitors):
    """
    Runs the full pipeline. Every stage is wrapped so that one failure
    (a bad scrape, a rate-limited call, malformed JSON) degrades gracefully
    instead of stopping the whole analysis — the user always ends up with
    whatever could be generated, plus clear warnings for what couldn't.
    """
    with st.status("Running competitive intelligence pipeline...", expanded=True) as status:

        # ---- Step 1: Discover competitors -----------------------------------
        st.write("🔎 Step 1/5 — Discovering competitors...")
        try:
            competitors = discover_competitors(business_domain, region, max_competitors)
        except Exception as e:
            st.warning(f"Competitor discovery hit an unexpected error: {e}")
            competitors = []

        if not competitors:
            status.update(label="No competitors found — try a broader domain/region.", state="error")
            return
        st.session_state.competitors = competitors

        # ---- Step 2: Scrape + analyze each competitor ------------------------
        st.write(f"🌐 Step 2/5 — Scraping & analyzing {len(competitors)} competitors...")
        scrape_progress = st.progress(0.0, text="Starting competitor analysis...")
        profiles = []
        for i, comp in enumerate(competitors):
            label = comp["name"] or comp["domain"]
            scrape_progress.progress(i / len(competitors), text=f"Analyzing {label}...")
            try:
                content = scrape_competitor_full(comp["url"])
            except Exception as e:
                st.warning(f"Could not scrape {label}: {e}")
                content = ""
            try:
                profile = analyze_competitor(llm, label, comp["url"], content)
            except Exception as e:
                st.warning(f"Analysis failed for {label}: {e}")
                profile = {}
            profiles.append(profile)
            time.sleep(0.2)
        scrape_progress.progress(1.0, text="Competitor analysis complete.")
        st.session_state.profiles = profiles

        # ---- Step 3: Market SWOT (Groq) --------------------------------------
        st.write("📊 Step 3/5 — Building market SWOT analysis (Groq)...")
        try:
            swot = build_swot(llm, business_domain, region, profiles)
        except Exception as e:
            st.warning(f"SWOT generation failed: {e}")
            swot = {"strengths": [], "weaknesses": [], "opportunities": [], "threats": [], "market_gaps": []}
        st.session_state.swot = swot

        # ---- Step 4: Growth strategy (Gemini — the one deep-reasoning call) --
        st.write("📈 Step 4/5 — Generating growth strategy & action plan (Gemini)...")
        try:
            strategy = build_growth_strategy(llm, business_domain, region, own_site, profiles, swot)
        except Exception as e:
            st.warning(f"Growth strategy generation failed: {e}")
            strategy = {
                "product_improvements": [], "new_features": [], "marketing_ideas": [],
                "seo_recommendations": [], "content_strategy": [],
                "customer_acquisition_ideas": [], "branding_suggestions": [],
                "expansion_opportunities": [], "plan_30_days": [], "plan_60_days": [],
                "plan_90_days": [],
            }
        st.session_state.strategy = strategy

        # ---- Step 5: Compile report -------------------------------------------
        st.write("📄 Step 5/5 — Compiling business intelligence report...")
        try:
            report_path = generate_report(
                business_domain, region, own_site, competitors, profiles, swot, strategy
            )
            st.session_state.report_path = report_path
        except Exception as e:
            st.warning(f"Report generation failed, but your analysis results are still available above: {e}")
            st.session_state.report_path = None

        status.update(label="Analysis complete!", state="complete")


def render_results(business_domain, region):
    competitors = st.session_state.competitors
    profiles = st.session_state.profiles
    swot = st.session_state.swot
    strategy = st.session_state.strategy

    if not competitors:
        st.info("Configure your business domain and region in the sidebar, then click **Run Competitive Analysis**.")
        return

    tabs = st.tabs(["🏢 Competitors", "📊 Comparison", "🧭 SWOT", "🚀 Growth Strategy", "📄 Report", "💬 Ask the AI"])

    with tabs[0]:
        for comp, profile in zip(competitors, profiles):
            with st.expander(f"**{comp['name'] or comp['domain']}** — {comp['domain']}", expanded=False):
                st.write(profile.get("company_overview", "Not available"))
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Products & Services**")
                    for p in profile.get("products_services") or []:
                        st.markdown(f"- {p}")
                    st.markdown(f"**Pricing:** {profile.get('pricing_info', 'Not available')}")
                    st.markdown(f"**Target Audience:** {profile.get('target_audience', 'Not available')}")
                with col2:
                    st.markdown("**Strengths**")
                    for s in profile.get("strengths") or []:
                        st.markdown(f"- ✅ {s}")
                    st.markdown("**Weaknesses**")
                    for w in profile.get("weaknesses") or []:
                        st.markdown(f"- ⚠️ {w}")

    with tabs[1]:
        rows = []
        for comp, profile in zip(competitors, profiles):
            rows.append({
                "Competitor": comp["name"] or comp["domain"],
                "Target Audience": profile.get("target_audience", "-"),
                "Pricing": profile.get("pricing_info", "-"),
                "Top Strength": (profile.get("strengths") or ["-"])[0],
                "Top Weakness": (profile.get("weaknesses") or ["-"])[0],
            })
        st.dataframe(rows, use_container_width=True)

    with tabs[2]:
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("### ✅ Strengths")
            for s in swot.get("strengths") or []:
                st.markdown(f"- {s}")
            st.markdown("### 🎯 Opportunities")
            for o in swot.get("opportunities") or []:
                st.markdown(f"- {o}")
        with c2:
            st.markdown("### ⚠️ Weaknesses")
            for w in swot.get("weaknesses") or []:
                st.markdown(f"- {w}")
            st.markdown("### 🚨 Threats")
            for t in swot.get("threats") or []:
                st.markdown(f"- {t}")
        st.markdown("### 🔍 Market Gaps")
        for g in swot.get("market_gaps") or []:
            st.markdown(f"- {g}")

    with tabs[3]:
        cols = st.columns(2)
        sections = [
            ("Product Improvements", "product_improvements"),
            ("New Features", "new_features"),
            ("Marketing Ideas", "marketing_ideas"),
            ("SEO Recommendations", "seo_recommendations"),
            ("Content Strategy", "content_strategy"),
            ("Customer Acquisition", "customer_acquisition_ideas"),
            ("Branding Suggestions", "branding_suggestions"),
            ("Expansion Opportunities", "expansion_opportunities"),
        ]
        for i, (label, key) in enumerate(sections):
            with cols[i % 2]:
                st.markdown(f"**{label}**")
                for item in strategy.get(key) or []:
                    st.markdown(f"- {item}")

        st.markdown("---")
        st.markdown("### 📅 30 / 60 / 90-Day Action Plan")
        p1, p2, p3 = st.columns(3)
        for col, label, key in [
            (p1, "First 30 Days", "plan_30_days"),
            (p2, "Next 60 Days", "plan_60_days"),
            (p3, "First 90 Days", "plan_90_days"),
        ]:
            with col:
                st.markdown(f"**{label}**")
                for item in strategy.get(key) or []:
                    st.markdown(f"- {item}")

    with tabs[4]:
        if st.session_state.report_path and os.path.exists(st.session_state.report_path):
            with open(st.session_state.report_path, "rb") as f:
                st.download_button(
                    "⬇️ Download Business Intelligence Report (.docx)",
                    data=f,
                    file_name="business_intelligence_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        else:
            st.info("Report not generated yet.")

    with tabs[5]:
        st.markdown("Ask a question about your competitive landscape (e.g. *\"What should my next business strategy be?\"*)")
        for role, msg in st.session_state.chat_history:
            with st.chat_message(role):
                st.write(msg)
        question = st.chat_input("Ask the AI business consultant...")
        if question:
            st.session_state.chat_history.append(("user", question))
            with st.chat_message("user"):
                st.write(question)
            llm = LLMClients(GROQ_API_KEY, GEMINI_API_KEY)
            context = {
                "business_domain": business_domain,
                "region": region,
                "competitors": competitors,
                "profiles": profiles,
                "swot": swot,
                "strategy": strategy,
            }
            try:
                answer = answer_business_question(llm, question, context)
            except Exception as e:
                st.warning(f"Chat request failed: {e}")
                answer = "Sorry, I couldn't process that question — please try again."
            st.session_state.chat_history.append(("assistant", answer))
            with st.chat_message("assistant"):
                st.write(answer)


def main():
    init_session_state()
    st.title("🚀 AI Business Strategy & Competitor Intelligence Agent")
    st.caption(
        "Automatically research competitors, extract insights, and generate a "
        "strategic growth report — powered by Gemini + Groq."
    )

    business_domain, region, own_site, max_competitors, run = sidebar_inputs()

    if run:
        if not business_domain or not region:
            st.error("Please provide both a business domain and a target region.")
        else:
            llm = LLMClients(GROQ_API_KEY, GEMINI_API_KEY)
            if not llm.is_ready():
                st.error(
                    "No working LLM client. Add GROQ_API_KEY and/or GEMINI_API_KEY "
                    "to your .env file and restart the app."
                )
            else:
                run_pipeline(llm, business_domain, region, own_site, max_competitors)

    render_results(business_domain, region)


if __name__ == "__main__":
    main()
